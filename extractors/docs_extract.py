from fastapi import FastAPI, UploadFile, File
from docx import Document
from PIL import Image
import pytesseract
import io
import uuid

app = FastAPI(title="DOC/DOCX Extractor")

def extract_docx(file_bytes: bytes):
    doc = Document(io.BytesIO(file_bytes))
    doc_id = f"doc_{uuid.uuid4().hex[:8]}"

    pages = []
    page_counter = 1

    # -------- TEXT BLOCKS --------
    text_blocks = []
    for para in doc.paragraphs:
        if para.text.strip():
            text_blocks.append({
                "block_id": f"text_p{page_counter}_{len(text_blocks)+1}",
                "text": para.text
            })

    # -------- TABLES --------
    tables = []
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])

        tables.append({
            "table_id": f"table_{t_idx+1}",
            "rows": rows
        })

    # -------- IMAGES (OCR ONLY – NO SAVE) --------
    images = []
    rels = doc.part._rels
    img_count = 1

    for rel in rels.values():
        if "image" in rel.reltype:
            image_bytes = rel.target_part.blob
            img = Image.open(io.BytesIO(image_bytes))
            ocr_text = pytesseract.image_to_string(img).strip()

            images.append({
                "image_id": f"{doc_id}_img{img_count}",
                "ocr_text": ocr_text
            })
            img_count += 1

    pages.append({
        "page_number": page_counter,
        "content": {
            "text_blocks": text_blocks,
            "tables": tables,
            "images": images
        }
    })

    return {
        "status": "success",
        "document": {
            "file_id": doc_id,
            "document_type": "docx",
            "total_pages": 1,
            "pages": pages
        }
    }


@app.post("/extract/doc")
async def extract_doc(file: UploadFile = File(...)):
    file_bytes = await file.read()
    return extract_docx(file_bytes)
